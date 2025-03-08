from flask import Flask, request, jsonify, render_template
from werkzeug.utils import secure_filename
import os
import PyPDF2
import pdfplumber
from docx import Document  # Import docx library
import google.generativeai as genai
from flask_cors import CORS
import zipfile


# Configure the app and allowed file types
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads/'
app.config['ALLOWED_EXTENSIONS'] = {'txt', 'pdf', 'docx'}

# Ensure the upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Configure API key for Google Generative AI (REPLACE WITH YOUR ACTUAL API KEY)
genai.configure(api_key="AIzaSyCpa3zWlaCsSoUIpy9u4YLZZBwxpJi-Aec") #<-- Replace with your actual key

# Enable CORS
CORS(app)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

@app.route('/')
def home():
    return render_template('index.html') #Corrected template name

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)

        try:
            file.save(filepath)
            extracted_text = extract_text(filepath, filename)
            #Removed AI cleaning as it was causing issues.  Consider re-adding with improved error handling if needed.
            return jsonify({"text": extracted_text})
        except Exception as e:
            return jsonify({"error": f"Error processing upload: {str(e)}"}), 500

    return jsonify({"error": "File type not allowed"}), 400

def extract_text(filepath, filename):
    ext = filename.split('.')[-1].lower()
    try:
        if ext == "txt":
            with open(filepath, "r", encoding="utf-8") as f:
                return f.read()
        elif ext == "pdf":
            return extract_text_from_pdf(filepath)
        elif ext == "docx":
            return extract_text_from_docx(filepath)
        else:
            return "Unsupported file format"
    except Exception as e:
        return f"Error reading file: {str(e)}"

def extract_text_from_pdf(filepath):
    text = ""
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
    except Exception:
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() or ""
    return text.strip() if text else "No readable text found in PDF"

def extract_text_from_docx(filepath):
    doc = Document(filepath)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)


def create_docx(text):
    try:
        document = Document()
        document.add_paragraph(text)
        temp_docx_path = "temp.docx"
        document.save(temp_docx_path)
        with open(temp_docx_path, "rb") as f:
            docx_data = f.read()
        os.remove(temp_docx_path)
        return docx_data
    except Exception as e:
        return f"Error creating DOCX: {str(e)}"



@app.route('/translate', methods=['POST'])
def translate_text():
    data = request.json
    text = data.get("text", "")
    target_language = data.get("targetLanguage", "English")
    output_structure = data.get("outputStructure", "formal and educational")

    if not text:
        return jsonify({"error": "Text is required"}), 400

    try:
        prompt = f"Translate the following text to {target_language} in a {output_structure} tone: '{text}'. " \
                 "Ensure that if the input is a single word, the output is a single word, and if the input is a sentence, " \
                 "the output retains the same meaning and structure. Provide only the translated text, nothing else."

        model = genai.GenerativeModel("gemini-2.0-flash")
        response = model.generate_content(prompt)
        translation = response.text if response else "Translation failed."
        return jsonify({"translation": translation})

    except genai.exception.GenerativeAIError as e:
        return jsonify({"error": f"Gemini API Error: {e}"}), 500
    except Exception as e:
        return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500


if __name__ == "__main__":
    app.run(debug=True)